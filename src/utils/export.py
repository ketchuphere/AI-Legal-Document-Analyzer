"""
src/utils/export.py
===================
Exports the final analysis report to DOCX format.
Returns bytes so Streamlit can offer a direct download button.
"""

import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def export_report_docx(result: dict) -> bytes:
    """
    Convert the analysis result dict into a formatted DOCX file.

    Args:
        result: The final state dict from the LangGraph workflow.

    Returns:
        DOCX file contents as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Title ────────────────────────────────────────────────────────
        title = doc.add_heading("AI Document Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  •  "
            f"Document: {result.get('doc_name', 'Unknown')}"
        ).runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph(
            "⚠️  Disclaimer: This report is AI-assisted and does not constitute legal advice. "
            "Consult a qualified attorney for professional guidance."
        ).runs[0].italic = True

        doc.add_paragraph("")

        # ── Executive Summary ────────────────────────────────────────────
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(result.get("summary", "Not available."))

        # ── Risk Analysis ────────────────────────────────────────────────
        doc.add_heading("Risk Analysis", level=1)
        risks = result.get("risks_structured", [])
        if risks:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
                "Risk", "Severity", "Likelihood", "Mitigation"
            )
            for r in risks:
                row = table.add_row().cells
                row[0].text = r.get("title", "—")
                row[1].text = r.get("severity", "—")
                row[2].text = r.get("likelihood", "—")
                row[3].text = r.get("mitigation", "—")
        else:
            doc.add_paragraph(result.get("risks", "Not available."))

        # ── Key Clauses ──────────────────────────────────────────────────
        doc.add_heading("Key Clause Extraction", level=1)
        clauses = result.get("clauses", {})
        if clauses:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Clause"
            table.rows[0].cells[1].text = "Extracted Value"
            for k, v in clauses.items():
                row = table.add_row().cells
                row[0].text = k.replace("_", " ").title()
                row[1].text = str(v) if v else "Not specified"
        else:
            doc.add_paragraph("No clause data available.")

        # ── Compliance ───────────────────────────────────────────────────
        doc.add_heading("Compliance Checklist", level=1)
        for item in result.get("compliance_items", []):
            status = item.get("status", "?")
            icon   = {"PASS": "✅", "FAIL": "❌", "WARN": "⚠️"}.get(status, "?")
            doc.add_paragraph(f"{icon} {item.get('check','')}: {item.get('note','')}")

        # ── Suggestions ──────────────────────────────────────────────────
        doc.add_heading("Improvement Suggestions", level=1)
        doc.add_paragraph(result.get("suggestions", "Not available."))

        # ── Serialise ────────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error("DOCX export failed: %s", e)
        # Return a minimal valid DOCX on error
        from docx import Document
        doc = Document()
        doc.add_paragraph("Export error: " + str(e))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
